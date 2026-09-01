from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


SOURCE = Path(r"E:\Download\绘语内网客户端在线更新版本设计与实施方案.docx")
OUTPUT = Path(r"E:\CodeX\linghui-im\docs\绘语内网客户端在线更新版本设计与实施方案-V1.1.docx")

NAVY = "17365D"
BLUE = "2F5597"
PALE_BLUE = "D9EAF7"
PALE_GREEN = "E2F0D9"
PALE_YELLOW = "FFF2CC"
PALE_RED = "FCE4D6"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"
TEXT = RGBColor(31, 45, 61)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def remove_document_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_font(run, name="Microsoft YaHei", size=10.5, bold=False, color=TEXT) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_repeat_header_footer(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Cm(2.1)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.0)
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.clear()
        run = paragraph.add_run("绘语内网在线更新 V1.1｜受控文档")
        set_font(run, size=8.5, color=RGBColor(100, 116, 139))
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.clear()
        run = paragraph.add_run("第 ")
        set_font(run, size=8.5, color=RGBColor(100, 116, 139))
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char1)
        run._r.append(instr_text)
        run._r.append(fld_char2)
        end = paragraph.add_run(" 页")
        set_font(end, size=8.5, color=RGBColor(100, 116, 139))


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in (("Title", 27, NAVY), ("Heading 1", 17, NAVY),
                              ("Heading 2", 13, BLUE), ("Heading 3", 11.5, BLUE)):
        style = document.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)


def add_paragraph(document: Document, text: str = "", *, bold_prefix: str | None = None,
                  style: str | None = None, align=None, color=None) -> None:
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_font(first, bold=True, color=color or TEXT)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_font(rest, color=color or TEXT)
    else:
        run = paragraph.add_run(text)
        set_font(run, color=color or TEXT)


def add_bullets(document: Document, items: list[str], level=0) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        set_font(run)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(item)
        set_font(run)


def add_callout(document: Document, title: str, text: str, fill=PALE_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 130, 170, 130, 170)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    first = paragraph.add_run(title + "　")
    set_font(first, bold=True, color=RGBColor.from_string(NAVY))
    rest = paragraph.add_run(text)
    set_font(rest)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    header = table.rows[0]
    repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        set_font(run, size=9, bold=True, color=RGBColor(255, 255, 255))
    for row_index, values in enumerate(rows):
        row = table.add_row()
        keep_row_together(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                set_cell_shading(cell, "F8FAFC")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_font(run, size=8.8)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_status_matrix(document: Document) -> None:
    add_table(document, ["工作项", "状态", "已落地证据", "下一步 / 验收"], [
        ["版本与发布基线", "已实现", "源码与锁文件已统一到 0.0.9；流水线校验标签、origin/main、HEAD 和干净工作区", "盘点现网版本、SHA-256、构建时间、设备数和 updater-capable 比例"],
        ["标签构建与签名门禁", "已实现", "标签推导 stable/beta；签名秘密仅进入构建步骤；Action 固定 SHA；校验安装包与 ArtTalk.exe", "在受审批 desktop-signing environment 完成一次真实签名构建"],
        ["不可变分发", "已实现", "$PublishRoot/<channel>/<version>/win-x64/；同版本异内容直接失败；HTTPS 全量回读", "用受信任内网证书验证 Nginx 实际回读与 Range 下载"],
        ["流水线草稿审批", "已实现", "独立 RELEASE_AUTOMATION_TOKEN 只能幂等创建/刷新 DRAFT；后台只改策略", "配置生产 Token，并验证发布/暂停权限无法被自动化 Token 获取"],
        ["发布状态与审计", "已实现", "DRAFT → PUBLISHED → PAUSED/REPLACED；原因、确认版本号与验证失败写专用审计表", "在预发布库执行迁移演练并导出审计证据"],
        ["客户端安全闭环", "已实现", "HTTPS 同源与不可变路径、releaseId/版本一致、安装前复核、暂停禁止安装、关闭退出自动安装", "使用两个已签名版本完成真实 E2E"],
        ["匿名定向与遥测", "已实现", "匿名遇 USER/DEPT 规则返回无更新；更新事件可匿名、有效 Token 再关联用户；Nginx 限流", "压测匿名事件限流并核对 releaseId 聚合"],
        ["专项自动化测试", "已实现", "桌面 40 项、管理端 4 项；后端策略/鉴权/迁移/产物回读用例；发布脚本负向用例", "补充生产证书错误、真实 404/摘要错误和完整安装重启 E2E"],
        ["现网基线盘点", "阻断", "当前环境无法确认员工电脑 0.0.3 是否含 updater，也无法访问现网地址", "盘点完成前不得填写最低支持版本或认定可直接在线升级"],
        ["正式发布验收", "待验证", "发布链路代码已闭环，但尚无受信任证书和两版签名安装包的现场证据", "按 0→10→30→100 灰度门槛执行并留存监控截图/日志"],
    ], widths=[1.25, 0.7, 2.75, 2.55])


def build_document() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document(SOURCE)
    remove_document_body(document)
    configure_styles(document)
    set_repeat_header_footer(document)
    document.core_properties.title = "绘语内网客户端在线更新版本设计与实施方案 V1.1"
    document.core_properties.subject = "可信构建、不可变分发、草稿审批和灰度发布闭环"
    document.core_properties.author = "绘语项目组"
    document.core_properties.comments = "由 V1.0 修订；原文档保持不变。"

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("绘语内网客户端在线更新\n版本设计与实施方案")
    set_font(run, size=27, bold=True, color=RGBColor.from_string(NAVY))
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("V1.1｜可信发布闭环修订版")
    set_font(run, size=16, bold=True, color=RGBColor.from_string(BLUE))
    document.add_paragraph()
    add_table(document, ["文档属性", "内容"], [
        ["修订日期", "2026-08-06"],
        ["候选版本", "0.0.9（不得覆盖 0.0.4～0.0.8 同版本产物）"],
        ["适用范围", "绘语 / ArtTalk Windows x64 内网客户端在线更新"],
        ["发布原则", "受控标签构建、可信签名、不可变 HTTPS 分发、流水线建草稿、管理员审批"],
        ["原文档", "《绘语内网客户端在线更新版本设计与实施方案》V1.0（保留不变）"],
    ], widths=[1.4, 5.7])
    add_callout(document, "发布前置结论", "现网 updater-capable 比例尚未盘点。0.0.9 先作为候选/引导版本；只有确认旧客户端含更新器后，才允许把 0.0.9 作为首次在线升级目标。", PALE_YELLOW)
    document.add_page_break()

    document.add_heading("1. 修订目标与决策基线", level=1)
    add_paragraph(document, "V1.1 将原 V1.0 的“功能建设计划”升级为可审计、可暂停、不可覆盖的正式发布闭环。原方案中的旧版本跳转、旧 HTTP 地址、共享更新清单、硬编码安装包名称和人工命令清单不再作为发布依据。")
    add_bullets(document, [
        "下一候选版本固定为 0.0.9；历史版本号一旦发布不得重构或覆盖。",
        "正式更新源统一为内网受信任 HTTPS；仅本机回环开发环境允许 HTTP。",
        "流水线只负责构建、签名、验证、不可变分发和创建后台草稿，不能发布或暂停。",
        "管理员只能编辑标题、日志、灰度、最低版本、强制更新和定向规则；核心产物字段永久只读。",
        "策略不可用、发布已暂停或 manifest 版本不一致时，客户端保持旧版本可用并禁止安装。",
    ])

    document.add_heading("2. 当前进度与阻断矩阵", level=1)
    add_status_matrix(document)

    document.add_heading("3. 真实发布基线", level=1)
    document.add_heading("3.1 现网盘点表", level=2)
    add_paragraph(document, "由桌面运维在发布前完成。任何未核实字段均记录为“未知”，不得用 V1.0 的 0.0.3 假设代替。")
    add_table(document, ["盘点项", "采集值", "证据", "判定"], [
        ["现网客户端版本分布", "待盘点", "设备清单 / 截图 / 资产系统导出", "阻断"],
        ["安装包 SHA-256 与构建时间", "待盘点", "安装介质哈希与文件元数据", "阻断"],
        ["“检查更新”入口与更新器能力", "待盘点", "抽样设备操作录像 / 日志", "阻断"],
        ["设备总数与 updater-capable 数量", "待盘点", "盘点表与去重设备 ID", "阻断"],
        ["内网 CA 是否被员工电脑信任", "待验证", "证书链检查与 HTTPS 请求结果", "发布前必须通过"],
    ], widths=[1.6, 1.0, 3.0, 1.4])
    document.add_heading("3.2 引导升级分支", level=2)
    add_numbered(document, [
        "若现网旧客户端确认包含可用更新器：以签名 0.0.9 完成首次在线升级验收。",
        "若不包含或无法确认：先人工部署签名 0.0.9 作为引导版本，再以签名 0.0.10 完成首次真实在线升级验收。",
        "首次正式发布保持 forceUpdate=false；minimumVersion 取盘点后确认的最旧 updater-capable 版本。",
    ])
    document.add_heading("3.3 发布一致性门禁", level=2)
    add_callout(document, "四项必须同时成立", "工作区干净；HEAD = origin/main；标签指向构建提交；标签版本 = package.json = package-lock.json。任一不一致，流水线失败。", PALE_RED)

    document.add_heading("4. 目标架构与状态模型", level=1)
    add_table(document, ["阶段", "责任主体", "输入", "输出 / 门禁"], [
        ["1. 标签", "发布负责人", "v0.0.9 或 v0.0.9-beta.1", "版本、通道、提交一致"],
        ["2. 构建签名", "受审批 Runner", "干净 origin/main", "签名 EXE、blockmap、manifest、provenance"],
        ["3. 校验分发", "发布脚本", "已签名产物", "不可变版本目录 + HTTPS 回读一致"],
        ["4. 草稿导入", "自动化 Token", "产物身份与验证证据", "DRAFT；不可发布/暂停"],
        ["5. 策略审批", "管理员", "标题、日志、灰度、定向、原因", "后端再次回读验证后 PUBLISHED"],
        ["6. 客户端执行", "Electron 客户端", "releaseId + 策略 + manifest", "检查、下载、安装前复核、VERSION_STARTED"],
    ], widths=[0.8, 1.15, 2.1, 3.1])
    add_paragraph(document, "实际状态模型：DRAFT → PUBLISHED → PAUSED / REPLACED。hasUpdate 不是人工开关，而是由 PUBLISHED 记录、客户端版本、定向规则和 rollout bucket 动态计算。")

    document.add_heading("5. 可信构建与不可变分发", level=1)
    document.add_heading("5.1 标签与通道", level=2)
    add_bullets(document, [
        "v0.0.9 对应 stable；v0.0.9-beta.1 对应 beta。其他预发布标识不进入 beta 流水线。",
        "第三方 GitHub Actions 固定到完整提交 SHA；签名作业使用受审批 desktop-signing environment。",
        "CSC_LINK 与 CSC_KEY_PASSWORD 仅注入“Build and sign NSIS package”步骤。",
    ])
    document.add_heading("5.2 产物与签名证据", level=2)
    add_table(document, ["对象", "命名 / 字段", "必须验证"], [
        ["安装包", "ArtTalk-Setup-<version>-x64.exe", "Valid 签名、发布者、证书指纹、时间戳、字节数、SHA-512"],
        ["应用主程序", "win-unpacked/ArtTalk.exe", "Valid 签名、同一发布者/证书指纹、可信时间戳"],
        ["更新清单", "latest.yml / beta.yml", "version、path/url、size、SHA-512 与安装包一致；SHA-256 留存"],
        ["差分元数据", "<installer>.blockmap", "存在、非空、HTTPS 回读内容一致"],
        ["构建证据", "release-provenance.json", "sourceCommit、产物大小/摘要、签名指纹、验证时间"],
    ], widths=[1.1, 2.4, 3.65])
    document.add_heading("5.3 不可变目录", level=2)
    add_callout(document, "唯一正式目录", "$PublishRoot/<channel>/<version>/win-x64/。目录通过同卷暂存后原子重命名就绪；已存在目录只有逐文件内容完全一致时才允许幂等重跑，同版本异内容必须失败。", PALE_GREEN)
    add_paragraph(document, "客户端策略中的 updateBaseUrl 必须为：https://<内网域名>/downloads/arttalk/<channel>/<version>/win-x64/。共享通道级清单不再使用。")

    document.add_heading("6. 流水线草稿审批闭环", level=1)
    add_numbered(document, [
        "发布脚本解析 manifest 得到真实安装包名称，不依赖硬编码旧文件名。",
        "暂存安装包、blockmap 和 manifest，校验版本、大小、SHA-512、签名发布者、证书指纹和时间戳。",
        "原子创建版本目录；通过内网 HTTPS 完整回读三类文件并再次比对 SHA-256。",
        "调用 POST /api/internal/client-release-drafts，使用独立 RELEASE_AUTOMATION_TOKEN 幂等创建或刷新 DRAFT。",
        "管理员填写策略与原因；发布时后端再次下载 manifest、EXE、blockmap 并校验摘要，验证成功才转为 PUBLISHED。",
    ])
    add_table(document, ["权限主体", "允许", "禁止"], [
        ["release-automation", "创建/刷新同一不可变身份的 DRAFT", "编辑策略、发布、暂停、替换"],
        ["管理员", "编辑策略、发布、暂停、查看统计", "编辑版本、URL、文件名、大小、摘要、提交、签名指纹"],
        ["匿名客户端", "查询公共策略、受限流上报设备事件", "访问管理接口、获取用户/部门定向结果"],
    ], widths=[1.3, 2.7, 3.15])

    document.add_heading("7. 接口、数据与审计", level=1)
    add_table(document, ["变更", "V1.1 约束"], [
        ["策略响应", "新增 releaseId；只返回有效 PUBLISHED 记录"],
        ["更新事件", "新增 releaseId；发布相关事件必须与版本/通道/平台/架构匹配；统计按 releaseId 聚合"],
        ["发布记录", "新增 sourceCommit、manifestName、manifestDigest、signerThumbprint、artifactVerifiedAt"],
        ["策略 PATCH", "PATCH /api/admin/client-releases/{id}/policy；仅策略字段可写；原因必填"],
        ["发布 / 暂停", "POST .../{id}/publish 与 .../{id}/pause；原因必填；强更或最低版本策略需输入确认版本号"],
        ["审计表", "记录 DRAFT_CREATED、POLICY_UPDATED、PUBLISHED、PAUSED、ARTIFACT_VERIFICATION_FAILED"],
        ["Flyway", "V20260806__harden_client_release_pipeline.sql；通过 information_schema 守卫兼容 schema.sql 与历史基线"],
    ], widths=[1.5, 5.65])

    document.add_heading("8. 客户端安全与可撤销性", level=1)
    add_bullets(document, [
        "serverOrigin 必须是 HTTPS origin；更新源必须同源，且 path 精确匹配 /downloads/arttalk/<channel>/<version>/win-x64/。",
        "策略 releaseId、策略版本和 manifest 版本必须一致；任何不一致均禁止下载/安装并上报 POLICY_MANIFEST_MISMATCH。",
        "autoInstallOnAppQuit=false。安装前重新请求策略，只有同一 releaseId 仍 hasUpdate 才允许 quitAndInstall。",
        "发布暂停、替换、策略不可达时，已下载包也不得安装；强制更新遮罩解除，旧版本继续可用。",
        "不再读取缓存强制策略；服务端不可用不触发离线锁屏。",
        "匿名请求遇 USER/DEPT 规则保守返回无更新；登录携带有效 Token 后重新评估。",
        "传输任务未完成时进入 WAIT_FOR_TRANSFERS；传输清零后仍必须重新执行策略复核。",
    ])

    document.add_heading("9. 测试与发布验收", level=1)
    document.add_heading("9.1 自动化覆盖", level=2)
    add_table(document, ["范围", "已覆盖", "现场补充"], [
        ["桌面更新器", "HTTPS/同源/不可变路径、stable/beta、release 身份、暂停/策略变化、传输安装门禁、错误状态", "真实 electron-updater + 两个签名版本"],
        ["后端", "0/10/30/100 灰度、ALLOW/DENY、匿名定向、草稿不可变、鉴权、状态转换、迁移资源、产物 HTTP 回读", "MySQL 现有库升级、空库初始化、并发发布压测"],
        ["管理端", "原因校验、强更版本二次确认、只读产物字段、构建", "浏览器交互与失败提示截图"],
        ["发布脚本", "恶意版本、缺失文件、摘要错误、错误签名、幂等重跑、同版本异内容、原子目录完整性", "内网 HTTPS、生产证书、真实共享目录权限"],
    ], widths=[1.2, 3.8, 2.15])
    document.add_heading("9.2 端到端场景", level=2)
    add_numbered(document, [
        "准备两个不同版本、同一可信发布者、带时间戳的签名安装包。",
        "通过受信任内网 HTTPS 完成检查、下载；下载过程中验证进度和 Range 请求。",
        "保持文件传输任务，确认安装等待；传输结束后再次请求策略。",
        "安装并重启，确认版本与 VERSION_STARTED releaseId 上报。",
        "分别注入错误签名、错误摘要、404、manifest 版本不一致、发布暂停，确认全部阻断且旧版本可用。",
        "执行消息、文件传输、托盘、登录和管理端核心业务回归。",
    ])
    document.add_heading("9.3 灰度节奏与门槛", level=2)
    add_table(document, ["阶段", "策略", "观察期", "推进条件"], [
        ["白名单", "3～5 台 DEVICE ALLOW，rollout=0", "2 小时", "签名/摘要零错误；核心功能零故障"],
        ["10%", "rollout=10", "4 小时", "下载成功率 ≥98%；测试批次启动成功率 ≥95%"],
        ["30%", "rollout=30", "1 个工作日", "失败率不升高；支持工单与业务指标正常"],
        ["100%", "rollout=100", "持续监控", "满足全部门槛并由发布负责人批准"],
    ], widths=[1.0, 2.4, 1.2, 2.6])
    add_callout(document, "立即暂停条件", "任一签名/摘要错误、核心业务回归、下载或安装失败率超过 2%。暂停后客户端安装前复核必须拒绝已下载包。", PALE_RED)

    document.add_heading("10. 发布操作规程（唯一入口）", level=1)
    add_paragraph(document, "本章取代 V1.0 中重复的人工复制和覆盖命令。正式发布只允许通过受控工作流和管理后台完成。")
    add_numbered(document, [
        "完成现网基线盘点，确认引导升级分支、最低支持版本和首批白名单。",
        "确保候选提交已推送 origin/main，package.json 与 package-lock.json 版本一致，评审通过后创建对应 v<semver> 标签。",
        "由受审批 desktop-signing environment 运行 Desktop release；核对 provenance、不可变目录和 HTTPS 回读结果。",
        "在管理后台打开流水线草稿，核对 sourceCommit、manifestDigest、signerThumbprint、artifactVerifiedAt 和安装包信息。",
        "填写更新标题、日志、DEVICE 白名单、rollout=0、forceUpdate=false、变更原因，保存策略。",
        "执行发布；后端产物复核成功后进入 PUBLISHED。按白名单 →10%→30%→100% 调整策略，每次填写原因并留存审计。",
        "达到暂停阈值立即执行暂停并记录事件编号、现象、影响和后续处置；禁止覆盖原版本目录。",
    ])

    document.add_heading("11. 证据清单与交付边界", level=1)
    add_table(document, ["证据", "位置 / 说明"], [
        ["标签发布工作流", ".github/workflows/desktop-release.yml"],
        ["不可变发布脚本与负向测试", "app/docker/scripts/publish-update.ps1；app/docker/scripts/tests/publish-update.tests.ps1"],
        ["客户端策略门禁", "app/im-web/electron/updater.ts；app/im-web/electron/update-policy.ts"],
        ["后端发布闭环", "ClientReleaseServiceImpl、ReleaseAutomationAuthenticationFilter、HttpClientReleaseArtifactVerifier"],
        ["数据库迁移", "V20260806__harden_client_release_pipeline.sql"],
        ["管理端审批", "app/im-admin/src/views/ClientReleaseManage.vue"],
        ["尚未形成的证据", "现网盘点、生产签名构建、内网 HTTPS 回读、两签名版本 E2E、生产数据库迁移演练"],
    ], widths=[2.0, 5.15])
    add_callout(document, "本次交付边界", "代码、测试与 V1.1 文档已修订；未创建 Git 标签、未推送、未签名或发布 0.0.9，也未修改现网策略。上述动作需由发布负责人基于盘点结果执行。", LIGHT_GRAY)

    document.add_heading("附录 A. 修订记录", level=1)
    add_table(document, ["版本", "日期", "修订摘要"], [
        ["V1.0", "原文档", "功能建设方案，使用旧版本跳转和旧 HTTP 地址基线"],
        ["V1.1", "2026-08-06", "切换到 0.0.9 候选、内网 HTTPS、标签构建、不可变目录、流水线草稿审批、releaseId、审计、安装前复核、灰度门槛与引导升级分支"],
    ], widths=[0.8, 1.2, 5.15])

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
