from pathlib import Path
from zipfile import ZipFile

from docx import Document


path = Path(r"E:\CodeX\linghui-im\docs\绘语内网客户端在线更新版本设计与实施方案-V1.1.docx")
with ZipFile(path) as archive:
    bad = archive.testzip()
    if bad:
        raise SystemExit(f"corrupt DOCX member: {bad}")

document = Document(path)
paragraph_text = [paragraph.text for paragraph in document.paragraphs]
table_text = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
text = "\n".join(paragraph_text + table_text)

required = [
    "V1.1", "0.0.9", "DRAFT → PUBLISHED → PAUSED / REPLACED",
    "$PublishRoot/<channel>/<version>/win-x64/", "POST /api/internal/client-release-drafts",
    "releaseId", "autoInstallOnAppQuit=false", "ARTIFACT_VERIFICATION_FAILED",
    "3～5 台 DEVICE ALLOW", "下载成功率 ≥98%", "现网盘点",
]
for value in required:
    if value not in text:
        raise SystemExit(f"required V1.1 content missing: {value}")

for value in ("http://172.16.59.253", "172.16.59.253:88", "ArtTalk Setup ", "stable/win-x64/latest.yml"):
    if value in text:
        raise SystemExit(f"stale V1.0 literal remains: {value}")

headings = [paragraph for paragraph in document.paragraphs if paragraph.style.name.startswith("Heading")]
if len(headings) < 15 or len(document.tables) < 15:
    raise SystemExit("document structure is unexpectedly incomplete")

print(f"DOCX structure OK: paragraphs={len(document.paragraphs)}, headings={len(headings)}, tables={len(document.tables)}")
